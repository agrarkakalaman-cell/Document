from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, font_size):
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    run.font.underline = False
    run.font.highlight_color = None

    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def set_page_number(section):
    section.different_first_page_header_footer = True

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()

    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)

    set_run_font(run, 14)


def format_paragraph(paragraph, font_size, line_spacing, first_line_indent):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.first_line_indent = Mm(first_line_indent)

    for run in paragraph.runs:
        set_run_font(run, font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.underline = False


def format_table(table, font_size, line_spacing):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = line_spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                for run in paragraph.runs:
                    set_run_font(run, font_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.underline = False


def format_dissertation(
    input_path,
    output_path,
    font_size,
    line_spacing,
    left_margin,
    right_margin,
    top_margin,
    bottom_margin,
    first_line_indent
):
    document = Document(input_path)

    styles = document.styles

    for style in styles:
        try:
            if style.type in [1, 2]:
                style.font.name = "Times New Roman"
                style.font.size = Pt(font_size)
                style.font.color.rgb = RGBColor(0, 0, 0)

                if style.element.rPr is not None:
                    style.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                    style.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    style.element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        except Exception:
            pass

    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        section.left_margin = Mm(left_margin)
        section.right_margin = Mm(right_margin)
        section.top_margin = Mm(top_margin)
        section.bottom_margin = Mm(bottom_margin)

        set_page_number(section)

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            format_paragraph(
                paragraph,
                font_size,
                line_spacing,
                first_line_indent
            )

    for table in document.tables:
        format_table(
            table,
            font_size,
            line_spacing
        )

    document.save(output_path)
