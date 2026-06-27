import streamlit as st
import streamlit.components.v1 as components
from pathlib import Path
from html import escape
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from utils.formatter import format_dissertation

st.set_page_config(page_title="Киберқауіпсіздік кафедрасы", layout="wide")

TEXTS = {
    "kk": {
        "page_title": "Құжат сапасын бақылау",
        "description": "Word құжатын жүктеңіз. Қосымша оны таңдалған рәсімдеу талаптарына сәйкес өңдейді.",
        "settings_title": "Талаптарды енгізіңіз:",
        "title_settings": "Титулка деректері",
        "document_type": "Құжат түрі",
        "topic": "Тақырыбы",
        "author": "Орындаған",
        "head": "Кафедра меңгерушісі",
        "advisor": "Ғылыми жетекші",
        "consultant": "Консультант (бөлім)",
        "control": "Нормоконтроль",
        "reviewer": "Рецензент",
        "code": "Шифр",
        "year": "Жыл",
        "font_size": "Қаріп өлшемі",
        "line_spacing": "Жоларалық интервал",
        "left_margin": "Сол жақ өріс, мм",
        "right_margin": "Оң жақ өріс, мм",
        "top_margin": "Жоғарғы өріс, мм",
        "bottom_margin": "Төменгі өріс, мм",
        "first_line_indent": "Абзац шегінісі, мм",
        "upload_button": "Жүктеу",
        "process_button": "Өңдеу",
        "clear_button": "Тазалау",
        "download_word": "Word жүктеп алу",
        "download_pdf": "PDF жүктеп алу",
        "file_uploader_label": "Word құжатын жүктеу",
        "changed_warning": "Талаптар өзгерді. Қайтадан «Өңдеу» батырмасын басыңыз.",
        "processing_error": "Құжатты өңдеу қатесі:",
        "template_missing": "Титулка шаблоны табылмады:",
        "language_label": "Тіл",
        "logo_missing": "logo.png файлы табылмады",
        "preview_title": "Алдын ала қарау",
        "preview_empty": "DOCX",
        "fill_required": "Титулка үшін Тақырыбы, Орындаған, Шифр, Жетекші және Жыл өрістерін толтырыңыз.",
        "doc_types": [
            "Дипломдық жұмыс",
            "Өндірістік практика есебі",
            "Оқу практикасы есебі",
            "Зертханалық жұмыс",
            "Курстық жұмыс",
            "СӨЖ",
            "СОӨЖ",
        ],
    },
    "ru": {
        "page_title": "Нормоконтроль документов",
        "description": "Загрузите Word-документ. Приложение приведет его к выбранным стандартам оформления.",
        "settings_title": "Введите требования:",
        "title_settings": "Данные титульного листа",
        "document_type": "Тип документа",
        "topic": "Тема",
        "author": "Выполнил(а)",
        "head": "Заведующий кафедрой",
        "advisor": "Руководитель",
        "consultant": "Консультанты (раздел)",
        "control": "Нормоконтроль",
        "reviewer": "Рецензент",
        "code": "Шифр",
        "year": "Год",
        "font_size": "Размер шрифта",
        "line_spacing": "Межстрочный интервал",
        "left_margin": "Левое поле, мм",
        "right_margin": "Правое поле, мм",
        "top_margin": "Верхнее поле, мм",
        "bottom_margin": "Нижнее поле, мм",
        "first_line_indent": "Абзацный отступ, мм",
        "upload_button": "Загрузить",
        "process_button": "Обработать",
        "clear_button": "Очистить",
        "download_word": "Скачать Word",
        "download_pdf": "Скачать PDF",
        "file_uploader_label": "Загрузить Word-документ",
        "changed_warning": "Требования изменились. Нажмите «Обработать» заново.",
        "processing_error": "Ошибка обработки документа:",
        "template_missing": "Шаблон титульного листа не найден:",
        "language_label": "Язык",
        "logo_missing": "Файл logo.png не найден",
        "preview_title": "Предварительный просмотр",
        "preview_empty": "DOCX",
        "fill_required": "Заполните поля: Тема, Выполнил(а), Шифр, Руководитель и Год.",
        "doc_types": [
            "Дипломная работа",
            "Отчет по производственной практике",
            "Отчет по учебной практике",
            "Лабораторная работа",
            "Курсовая работа",
            "СРО",
            "СРСП",
        ],
    },
}

TEMPLATE_MAP = {
    "Дипломдық жұмыс": "templates/Diploma_kaz.docx",
    "Дипломная работа": "templates/Diploma_rus.docx",
    "Өндірістік практика есебі": "templates/Practic_kaz.docx",
    "Отчет по производственной практике": "templates/Practic_rus.docx",
    "Оқу практикасы есебі": "templates/Study_practic_kaz.docx",
    "Отчет по учебной практике": "templates/Study_practic_rus.docx",
    "Зертханалық жұмыс": "templates/Lab_kaz.docx",
    "Лабораторная работа": "templates/Lab_rus.docx",
    "Курстық жұмыс": "templates/Course_kaz.docx",
    "Курсовая работа": "templates/Course_rus.docx",
    "СӨЖ": "templates/SRO_kaz.docx",
    "СРО": "templates/SRO_rus.docx",
    "СОӨЖ": "templates/SRSP_kaz.docx",
    "СРСП": "templates/SRSP_rus.docx",
}

DEFAULT_SETTINGS = {
    "font_size": 14,
    "line_spacing": 1.5,
    "left_margin": 30,
    "right_margin": 20,
    "top_margin": 20,
    "bottom_margin": 20,
    "first_line_indent": 12.5,
}

if "lang" not in st.session_state:
    st.session_state["lang"] = "kk"

current_lang = st.session_state["lang"]
t = TEXTS[current_lang]

SESSION_DEFAULTS = {
    "docx_path": None,
    "error_message": None,
    "uploader_key": 0,
    "last_settings": None,
    "processed_settings": None,
    "preview_html": None,
    "document_type": t["doc_types"][0],
    "title_topic": "",
    "title_author": "",
    "title_head": "",
    "title_advisor": "",
    "title_consultants": "",
    "title_control": "",
    "title_reviewer": "",
    "title_code": "",
    "title_year": str(datetime.now().year),
}

for key, value in SESSION_DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value

for key, value in DEFAULT_SETTINGS.items():
    if key not in st.session_state:
        st.session_state[key] = value


def clear_all():
    st.session_state["docx_path"] = None
    st.session_state["error_message"] = None
    st.session_state["processed_settings"] = None
    st.session_state["last_settings"] = None
    st.session_state["preview_html"] = None
    st.session_state["uploader_key"] += 1

    for key, value in DEFAULT_SETTINGS.items():
        st.session_state[key] = value

    st.session_state["document_type"] = t["doc_types"][0]
    st.session_state["title_topic"] = ""
    st.session_state["title_author"] = ""
    st.session_state["title_head"] = ""
    st.session_state["title_advisor"] = ""
    st.session_state["title_consultants"] = ""
    st.session_state["title_control"] = ""
    st.session_state["title_reviewer"] = ""
    st.session_state["title_code"] = ""
    st.session_state["title_year"] = str(datetime.now().year)

    for folder in [Path("uploads"), Path("outputs")]:
        if folder.exists():
            for file_path in folder.glob("*"):
                if file_path.is_file():
                    try:
                        file_path.unlink()
                    except OSError:
                        pass


def replace_placeholders_in_docx(doc: Document, replacements: dict):
    def replace_in_paragraph(paragraph):
        text = paragraph.text
        if not any(key in text for key in replacements):
            return
        for key, value in replacements.items():
            text = text.replace(key, str(value))
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def append_docx_body(target_doc: Document, source_doc: Document):
    target_body = target_doc.element.body
    source_body = source_doc.element.body

    target_sect_pr = None
    if len(target_body) and target_body[-1].tag.endswith("sectPr"):
        target_sect_pr = target_body[-1]
        target_body.remove(target_sect_pr)

    p = target_doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    for element in source_body:
        if element.tag.endswith("sectPr"):
            continue
        target_body.append(element)

    if target_sect_pr is not None:
        target_body.append(target_sect_pr)


def create_document_with_title_page(template_path, formatted_docx_path, output_docx_path, data):
    template_path = Path(template_path)
    formatted_docx_path = Path(formatted_docx_path)
    output_docx_path = Path(output_docx_path)

    if not template_path.exists():
        raise FileNotFoundError(f"{t['template_missing']} {template_path}")

    title_doc = Document(template_path)
    replace_placeholders_in_docx(title_doc, data)

    main_doc = Document(formatted_docx_path)
    append_docx_body(title_doc, main_doc)

    output_docx_path.parent.mkdir(exist_ok=True)
    title_doc.save(output_docx_path)


def get_docx_preview_html(docx_path, max_paragraphs=220):
    document = Document(docx_path)
    html_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text
        if not text.strip():
            html_parts.append("<p>&nbsp;</p>")
        else:
            alignment = "left"
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = "center"
            elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = "right"
            elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = "justify"
            html_parts.append(f'<p style="text-align: {alignment};">{escape(text)}</p>')

        if len(html_parts) >= max_paragraphs:
            break

    return "".join(html_parts)


def show_docx_preview():
    if st.session_state.get("preview_html"):
        st.markdown(
            f'<div class="docx-preview">{st.session_state["preview_html"]}</div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f'<div class="docx-preview docx-preview-empty">{t["preview_empty"]}</div>',
            unsafe_allow_html=True
        )


def print_pdf_button(preview_html, button_text, disabled=False):
    if disabled:
        st.button(button_text, disabled=True, use_container_width=True)
        return

    safe_html = preview_html.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")

    components.html(
        f"""
        <button class="pdf-print-btn" onclick="printDocxPreview()">{button_text}</button>
        <script>
        function printDocxPreview() {{
            const html = `{safe_html}`;
            const w = window.open("", "_blank");
            w.document.write(`
                <html>
                <head>
                    <title>PDF</title>
                    <style>
                        @page {{ size: A4; margin: {st.session_state["top_margin"]}mm {st.session_state["right_margin"]}mm {st.session_state["bottom_margin"]}mm {st.session_state["left_margin"]}mm; }}
                        body {{
                            font-family: "Times New Roman", serif;
                            font-size: {st.session_state["font_size"]}pt;
                            line-height: {st.session_state["line_spacing"]};
                            color: #000;
                        }}
                        p {{ margin: 0 0 10px 0; }}
                    </style>
                </head>
                <body>${{html}}</body>
                </html>
            `);
            w.document.close();
            w.focus();
            setTimeout(() => w.print(), 300);
        }}
        </script>
        <style>
            .pdf-print-btn {{
                width: 100%; height: 39px; background: #eaf3ff; color: black; border: 1px solid #c8d3e1;
                border-radius: 7px; font-size: 14px; cursor: pointer; font-family: sans-serif;
            }}
        </style>
        """,
        height=46
    )


# ── CSS Стильдер ──────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
.block-container {{ padding-top: 34px; padding-left: 34px; padding-right: 34px; max-width: 1400px; }}
.header-title {{ font-size: 24px; font-weight: 700; text-align: center; margin-top: 20px; margin-bottom: 14px; }}
.description-text {{ font-size: 16px; margin-bottom: 34px; }}
.col-title {{ font-size: 16px; font-weight: 700; margin-top: 0px; margin-bottom: 14px; color: #1f2937; }}
.docx-preview {{
    width: 100%; height: 750px; border: 5px solid #f1f3f6; background-color: white; box-sizing: border-box;
    margin-top: 12px; margin-bottom: 16px;
    padding-top: {DEFAULT_SETTINGS["top_margin"]}mm; padding-right: {DEFAULT_SETTINGS["right_margin"]}mm;
    padding-bottom: {DEFAULT_SETTINGS["bottom_margin"]}mm; padding-left: {DEFAULT_SETTINGS["left_margin"]}mm;
    overflow-y: auto; font-family: "Times New Roman", serif; font-size: {DEFAULT_SETTINGS["font_size"]}pt;
    line-height: {DEFAULT_SETTINGS["line_spacing"]}; color: black; white-space: normal;
}}
.docx-preview-empty {{ display: flex; align-items: center; justify-content: center; color: #eef0f4; font-size: 90px; font-weight: 700; padding: 0; }}
div[data-testid="stNumberInput"] {{ margin-bottom: 14px; }}
div[data-testid="stTextInput"] {{ margin-bottom: 10px; }}
div[data-testid="stNumberInput"] label, div[data-testid="stTextInput"] label, div[data-testid="stSelectbox"] label {{ font-size: 13px; font-weight: 500; }}
div[data-testid="stNumberInput"] input, div[data-testid="stTextInput"] input {{ background-color: #eaf3ff; }}
div[data-testid="stFileUploader"] {{ width: 100%; margin-bottom: 10px; }}
div[data-testid="stFileUploader"] section {{ width: 100%; padding: 0 !important; border: 1px solid #c8d3e1 !important; border-radius: 7px !important; background: #eaf3ff !important; min-height: 46px !important; }}
div[data-testid="stFileUploaderDropzone"] {{ width: 100%; padding: 0 !important; border: none !important; background: #eaf3ff !important; display: flex; align-items: center; justify-content: center; min-height: 44px !important; }}
[data-testid="stFileUploaderDropzoneInstructions"], [data-testid="stFileUploaderFile"], [data-testid="stFileUploaderFileData"], div[data-testid="stFileUploader"] small, [data-testid="stFileUploader"] section small, [data-testid="stFileUploader"] div small {{ display: none !important; }}
div[data-testid="stFileUploader"] button {{ width: 100% !important; height: 44px !important; min-width: 100% !important; background: #eaf3ff !important; border: none !important; color: transparent !important; box-shadow: none !important; position: relative; margin: 0 !important; }}
div[data-testid="stFileUploader"] button::after {{ color: black; font-size: 14px; font-weight: 600; position: absolute; left: 50%; top: 50%; transform: translate(-50%, -50%); white-space: nowrap; }}
.stButton button, .stDownloadButton button {{ width: 100% !important; height: 39px; background-color: #eaf3ff !important; color: black !important; border: 1px solid #c8d3e1 !important; border-radius: 7px !important; font-weight: 400 !important; margin: 0 !important; box-shadow: none !important; }}
.stButton button:hover, .stDownloadButton button:hover {{ border-color: #9fb5cc !important; }}
@media screen and (max-width: 800px) {{ .block-container {{ padding-left: 14px; padding-right: 14px; }} .docx-preview {{ height: 420px; }} .docx-preview-empty {{ font-size: 64px; }} }}
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
top_col_1, top_col_2, top_col_3 = st.columns([1, 4, 1], gap="medium")

with top_col_1:
    logo_path = Path("logo.png")
    if logo_path.exists():
        st.markdown("<div style='margin-top:20px;'></div>", unsafe_allow_html=True)
        st.image(str(logo_path), width=180)
    else:
        st.warning(t["logo_missing"])

with top_col_2:
    st.markdown(f'<div class="header-title">{t["page_title"]}</div>', unsafe_allow_html=True)

with top_col_3:
    st.markdown('<div style="height:26px;"></div>', unsafe_allow_html=True)
    selected_language = st.selectbox(
        t["language_label"],
        options=["Қазақша", "Русский"],
        index=0 if st.session_state["lang"] == "kk" else 1,
        label_visibility="collapsed"
    )
    new_lang = "kk" if selected_language == "Қазақша" else "ru"
    if new_lang != st.session_state["lang"]:
        st.session_state["lang"] = new_lang
        st.session_state["document_type"] = TEXTS[new_lang]["doc_types"][0]
        st.rerun()

st.markdown(f'<div class="description-text">{t["description"]}</div>', unsafe_allow_html=True)

input_dir = Path("uploads")
output_dir = Path("outputs")
input_dir.mkdir(exist_ok=True)
output_dir.mkdir(exist_ok=True)

# ── 3 негізгі баған ───────────────────────────────────────────────────────────
left_col, mid_col, right_col = st.columns([1.2, 1.8, 1.2], gap="large")

# ── СОЛ ЖАҚ: Титулка деректері ───────────────────────────────────────────────
with left_col:
    st.markdown(f'<div class="col-title">{t["title_settings"]}</div>', unsafe_allow_html=True)

    try:
        default_index = t["doc_types"].index(st.session_state["document_type"])
    except ValueError:
        default_index = 0

    document_type     = st.selectbox(t["document_type"], options=t["doc_types"], index=default_index, key="document_type")
    title_topic       = st.text_input(t["topic"],       key="title_topic")
    title_author      = st.text_input(t["author"],      key="title_author")
    title_advisor     = st.text_input(t["advisor"],     key="title_advisor")
    title_head        = st.text_input(t["head"],        key="title_head")
    title_consultants = st.text_input(t["consultant"],  key="title_consultants")
    title_control     = st.text_input(t["control"],     key="title_control")
    title_reviewer    = st.text_input(t["reviewer"],    key="title_reviewer")
    title_code        = st.text_input(t["code"],        key="title_code")
    title_year        = st.text_input(t["year"],        key="title_year")

# ── ОРТА: Файл жүктеу, батырмалар, preview, жүктеп алу ──────────────────────
with mid_col:
    uploaded_file = st.file_uploader(
        t["file_uploader_label"],
        type=["docx"],
        label_visibility="collapsed",
        key=f"uploader_{st.session_state['uploader_key']}"
    )

    if uploaded_file is None:
        file_status = t["upload_button"]
    else:
        file_status = f"📄 {uploaded_file.name}   {uploaded_file.size / 1024:.1f} KB"

    file_status_escaped = file_status.replace("\\", "\\\\").replace('"', '\\"')
    st.markdown(
        f"<style>div[data-testid='stFileUploader'] button::after {{ content: '{file_status_escaped}'; }}</style>",
        unsafe_allow_html=True
    )

    button_col_1, button_col_2 = st.columns([1, 1], gap="medium")

    input_path, formatted_docx_path, final_docx_path = None, None, None
    if uploaded_file is not None:
        input_path = input_dir / uploaded_file.name
        file_stem = Path(uploaded_file.name).stem
        formatted_docx_path = output_dir / f"formatted_{file_stem}.docx"
        final_docx_path = output_dir / f"final_{file_stem}.docx"
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

    with button_col_1:
        process_clicked = st.button(t["process_button"], disabled=uploaded_file is None, use_container_width=True)

    with button_col_2:
        st.button(t["clear_button"], use_container_width=True, on_click=clear_all)

    if process_clicked and uploaded_file is not None:
        st.session_state["docx_path"] = None
        st.session_state["error_message"] = None
        st.session_state["preview_html"] = None

        # right_col виджеттерінің мәндері session_state-тен алынады
        _font_size         = st.session_state.get("font_size",         DEFAULT_SETTINGS["font_size"])
        _line_spacing      = st.session_state.get("line_spacing",      DEFAULT_SETTINGS["line_spacing"])
        _left_margin       = st.session_state.get("left_margin",       DEFAULT_SETTINGS["left_margin"])
        _right_margin      = st.session_state.get("right_margin",      DEFAULT_SETTINGS["right_margin"])
        _top_margin        = st.session_state.get("top_margin",        DEFAULT_SETTINGS["top_margin"])
        _bottom_margin     = st.session_state.get("bottom_margin",     DEFAULT_SETTINGS["bottom_margin"])
        _first_line_indent = st.session_state.get("first_line_indent", DEFAULT_SETTINGS["first_line_indent"])

        if not all([title_topic.strip(), title_author.strip(), title_advisor.strip(), title_code.strip(), title_year.strip()]):
            st.session_state["error_message"] = t["fill_required"]
        else:
            try:
                format_dissertation(
                    input_path=input_path, output_path=formatted_docx_path,
                    font_size=_font_size, line_spacing=_line_spacing,
                    left_margin=_left_margin, right_margin=_right_margin,
                    top_margin=_top_margin, bottom_margin=_bottom_margin,
                    first_line_indent=_first_line_indent
                )

                if not formatted_docx_path.exists():
                    st.session_state["error_message"] = t["processing_error"]
                else:
                    template_path = TEMPLATE_MAP.get(document_type)
                    if not template_path:
                        raise FileNotFoundError(f"Template not mapped for {document_type}")

                    replacements = {
                        "{{DOC_TYPE}}":   document_type,
                        "{{TITLE}}":      title_topic,
                        "{{AUTHOR}}":     title_author,
                        "{{ADVISOR}}":    title_advisor,
                        "{{HEAD}}":       title_head,
                        "{{CONSULTANT}}": title_consultants,
                        "{{CONTROL}}":    title_control,
                        "{{REVIEWER}}":   title_reviewer,
                        "{{CODE}}":       title_code,
                        "{{YEAR}}":       title_year,
                    }

                    create_document_with_title_page(
                        template_path=template_path,
                        formatted_docx_path=formatted_docx_path,
                        output_docx_path=final_docx_path,
                        data=replacements
                    )

                    st.session_state["docx_path"] = str(final_docx_path)
                    st.session_state["preview_html"] = get_docx_preview_html(final_docx_path)
                    st.session_state["processed_settings"] = {
                        "document_type":    document_type,
                        "title_topic":      title_topic,
                        "title_author":     title_author,
                        "title_advisor":    title_advisor,
                        "title_head":       title_head,
                        "title_consultants": title_consultants,
                        "title_control":    title_control,
                        "title_reviewer":   title_reviewer,
                        "title_code":       title_code,
                        "title_year":       title_year,
                        "font_size":         _font_size,
                        "line_spacing":      _line_spacing,
                        "left_margin":       _left_margin,
                        "right_margin":      _right_margin,
                        "top_margin":        _top_margin,
                        "bottom_margin":     _bottom_margin,
                        "first_line_indent": _first_line_indent,
                    }
                    st.rerun()

            except Exception as error:
                st.session_state["error_message"] = f'{t["processing_error"]} {error}'

    if st.session_state["error_message"]:
        st.error(st.session_state["error_message"])

    st.markdown(f'<div style="font-weight:700; text-align:center; margin-top:12px;">{t["preview_title"]}</div>', unsafe_allow_html=True)
    show_docx_preview()

    download_col_1, download_col_2 = st.columns([1, 1], gap="medium")
    with download_col_1:
        if st.session_state["docx_path"]:
            ready_docx = Path(st.session_state["docx_path"])
            with open(ready_docx, "rb") as f:
                st.download_button(
                    label=t["download_word"], data=f.read(),
                    file_name=ready_docx.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.button(t["download_word"], disabled=True, use_container_width=True)

    with download_col_2:
        print_pdf_button(
            st.session_state.get("preview_html") or "",
            t["download_pdf"],
            disabled=not bool(st.session_state.get("preview_html"))
        )

# ── ОҢ ЖАҚ: Сандық талаптар ──────────────────────────────────────────────────
with right_col:
    st.markdown(f'<div class="col-title">{t["settings_title"]}</div>', unsafe_allow_html=True)

    font_size         = st.number_input(t["font_size"],         min_value=8,   max_value=30,  value=st.session_state["font_size"],          step=1,   key="font_size")
    line_spacing      = st.number_input(t["line_spacing"],      min_value=1.0, max_value=3.0, value=st.session_state["line_spacing"],        step=0.1, format="%.2f", key="line_spacing")
    left_margin       = st.number_input(t["left_margin"],       min_value=0,   max_value=50,  value=st.session_state["left_margin"],         step=1,   key="left_margin")
    right_margin      = st.number_input(t["right_margin"],      min_value=0,   max_value=50,  value=st.session_state["right_margin"],        step=1,   key="right_margin")
    top_margin        = st.number_input(t["top_margin"],        min_value=0,   max_value=50,  value=st.session_state["top_margin"],          step=1,   key="top_margin")
    bottom_margin     = st.number_input(t["bottom_margin"],     min_value=0,   max_value=50,  value=st.session_state["bottom_margin"],       step=1,   key="bottom_margin")
    first_line_indent = st.number_input(t["first_line_indent"], min_value=0.0, max_value=30.0, value=st.session_state["first_line_indent"],  step=0.5, format="%.2f", key="first_line_indent")

# ── «Талаптар өзгерді» ескертуі ──────────────────────────────────────────────
current_settings = {
    "document_type":    st.session_state.get("document_type"),
    "title_topic":      st.session_state.get("title_topic"),
    "title_author":     st.session_state.get("title_author"),
    "title_advisor":    st.session_state.get("title_advisor"),
    "title_head":       st.session_state.get("title_head"),
    "title_consultants": st.session_state.get("title_consultants"),
    "title_control":    st.session_state.get("title_control"),
    "title_reviewer":   st.session_state.get("title_reviewer"),
    "title_code":       st.session_state.get("title_code"),
    "title_year":       st.session_state.get("title_year"),
    "font_size":         font_size,
    "line_spacing":      line_spacing,
    "left_margin":       left_margin,
    "right_margin":      right_margin,
    "top_margin":        top_margin,
    "bottom_margin":     bottom_margin,
    "first_line_indent": first_line_indent,
}

if st.session_state["docx_path"] and st.session_state["processed_settings"] != current_settings:
    with mid_col:
        st.warning(t["changed_warning"])
